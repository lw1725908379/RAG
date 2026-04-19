# -*- coding: utf-8 -*-
"""
AI核心层 - 文档解析模块
支持：JSON、TXT、Markdown、Word、PDF
"""
import os
import json
import logging
from typing import List, Dict, Any, Optional

logger = logging.getLogger(__name__)

# 尝试导入可选库
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 未安装，Word文件解析将不可用")

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pdfplumber 未安装，PDF文件解析将不可用")


class DocumentParser:
    """
    文档解析器
    支持多种格式转换为统一文档格式
    """

    @staticmethod
    def parse(file_path: str) -> List[Dict[str, Any]]:
        """
        解析文档

        Args:
            file_path: 文件路径

        Returns:
            文档列表 [{"id": "...", "title": "...", "content": "..."}]
        """
        ext = os.path.splitext(file_path)[1].lower()

        parsers = {
            '.json': DocumentParser._parse_json,
            '.txt': DocumentParser._parse_txt,
            '.md': DocumentParser._parse_md,
            '.docx': DocumentParser._parse_docx,
            '.doc': DocumentParser._parse_docx,
            '.pdf': DocumentParser._parse_pdf,
        }

        parser = parsers.get(ext)
        if not parser:
            raise ValueError(f"不支持的文件格式: {ext}")

        return parser(file_path)

    @staticmethod
    def _parse_json(file_path: str) -> List[Dict[str, Any]]:
        """解析JSON文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # 支持两种格式：直接是列表，或 {"data": [...]} 或 {"cases": [...]}
        if isinstance(data, list):
            items = data
        elif isinstance(data, dict):
            items = data.get('data', data.get('cases', []))
        else:
            raise ValueError("JSON格式不支持")

        results = []
        for i, item in enumerate(items):
            # 兼容不同字段名
            title = item.get('title', item.get('名称', item.get('name', f'文档{i+1}')))
            content = item.get('content', item.get('内容', item.get('description', item.get('描述', ''))))
            doc_id = item.get('id', item.get('case_id', f'json_{i+1}'))

            if title or content:
                results.append({
                    "id": str(doc_id),
                    "title": str(title) if title else '',
                    "content": str(content) if content else ''
                })

        return results

    @staticmethod
    def _parse_txt(file_path: str) -> List[Dict[str, Any]]:
        """解析TXT文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 按空行分隔段落
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]

        results = []
        for i, para in enumerate(paragraphs):
            # 取第一行作为标题
            lines = para.split('\n')
            title = lines[0].strip() if lines else f'段落{i+1}'
            content = '\n'.join(lines[1:]).strip() if len(lines) > 1 else para

            results.append({
                "id": f"txt_{i+1}",
                "title": title[:100],
                "content": content
            })

        return results

    @staticmethod
    def _parse_md(file_path: str) -> List[Dict[str, Any]]:
        """解析Markdown文件 - 按语义分块"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        import re

        # 移除目录部分
        content = re.sub(r'^#*\s*目[录录].*?(?=#+\s|\Z)', '', content, flags=re.DOTALL | re.IGNORECASE)

        # 按章节分隔
        sections = re.split(r'\n(?=#+\s)', content)

        results = []
        chunk_id = 1
        overlap_tokens = 100  # 重叠约100个字符(约50-80 tokens)

        for i, section in enumerate(sections):
            section = section.strip()
            if not section:
                continue

            # 提取标题
            match = re.match(r'^(#{1,6})\s+(.+?)\n', section)
            if match:
                title = match.group(2).strip()
                body = section[match.end():].strip()
            else:
                title = f'章节{i+1}'
                body = section

            # 跳过目录
            if re.search(r'^目[录录]|^table\s+of', title, re.IGNORECASE):
                continue

            # 按段落分块（以空行或换行分隔）
            paragraphs = re.split(r'\n\s*\n', body)

            current_chunk = ""
            prev_tail = ""  # 上一段结尾（用于重叠）

            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                # 计算当前块大小
                new_content = current_chunk + "\n\n" + para if current_chunk else para

                # 如果超过阈值，分块
                if len(new_content) > 500:
                    # 保存当前块
                    if current_chunk:
                        results.append({
                            "id": f"chunk_{chunk_id}",
                            "title": title[:80],
                            "content": current_chunk
                        })
                        chunk_id += 1

                    # 保留尾部用于重叠
                    tail_len = min(len(current_chunk), overlap_tokens)
                    prev_tail = current_chunk[-tail_len:] if tail_len > 0 else ""
                    current_chunk = prev_tail + para
                else:
                    current_chunk = new_content

            # 保存最后一块
            if current_chunk:
                effective = re.sub(r'[\s<][^>]*>', '', current_chunk)
                if len(effective.strip()) >= 20:
                    results.append({
                        "id": f"chunk_{chunk_id}",
                        "title": title[:80],
                        "content": current_chunk
                    })

        return results

    @staticmethod
    def _parse_docx(file_path: str) -> List[Dict[str, Any]]:
        """解析Word文档 (.docx)"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")

        doc = docx.Document(file_path)

        results = []
        current_title = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 判断是否是标题 (样式以 Heading 开头)
            if para.style.name.startswith('Heading'):
                # 保存之前的段落
                if current_content:
                    results.append({
                        "id": f"docx_{len(results)+1}",
                        "title": current_title or f'段落{len(results)+1}',
                        "content": '\n'.join(current_content)
                    })

                current_title = text
                current_content = []
            else:
                current_content.append(text)

        # 保存最后一段
        if current_content:
            results.append({
                "id": f"docx_{len(results)+1}",
                "title": current_title or f'段落{len(results)+1}',
                "content": '\n'.join(current_content)
            })

        # 如果没有标题，检查表格
        if not results:
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 2:
                        results.append({
                            "id": f"table_{i+1}",
                            "title": cells[0][:100],
                            "content": cells[1]
                        })

        return results if results else [{"id": "docx_1", "title": "文档", "content": "无法解析文档内容"}]

    @staticmethod
    def _parse_pdf(file_path: str) -> List[Dict[str, Any]]:
        """解析PDF文件"""
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber 未安装，请运行: pip install pdfplumber")

        results = []

        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # 提取文本
                text = page.extract_text()
                if not text:
                    continue

                # 按行分组
                lines = [l.strip() for l in text.split('\n') if l.strip()]

                # 尝试识别标题
                title = None
                content_lines = []

                for line in lines:
                    # 大写或短行可能是标题
                    if len(line) < 50 and (line.isupper() or line.endswith(':')):
                        if content_lines:
                            # 保存前一段
                            results.append({
                                "id": f"pdf_p{page_num}_{len(results)+1}",
                                "title": title or f'第{page_num}页',
                                "content": '\n'.join(content_lines)
                            })
                            content_lines = []

                        title = line
                    else:
                        content_lines.append(line)

                # 保存最后一段
                if content_lines:
                    results.append({
                        "id": f"pdf_p{page_num}_{len(results)+1}",
                        "title": title or f'第{page_num}页',
                        "content": '\n'.join(content_lines)
                    })

        return results if results else [{"id": "pdf_1", "title": "PDF文档", "content": "无法解析PDF内容"}]


def parse_documents(file_path: str) -> List[Dict[str, Any]]:
    """
    便捷函数：解析文档

    Args:
        file_path: 文件路径

    Returns:
        文档列表
    """
    return DocumentParser.parse(file_path)
